"""
app/capabilities/documents/ingestion.py — Document parsing, chunking & embedding pipeline.

Flow (runs in a background task after upload):
  1. parse_document()     — extract text from the raw bytes based on MIME type
  2. chunk_text()         — split into overlapping token-sized windows
  3. embed_and_store()    — embed each chunk + bulk-upsert into document_chunks

Supported formats (Phase 1):
  • PDF   — pypdf text extraction (text-native only; scanned PDFs return empty pages)
  • DOCX  — python-docx paragraph extraction
  • TXT / MD — raw UTF-8 decode
  • CSV   — formatted as "col: val, col: val" rows for semantic search

OCR for scanned PDFs is planned for Phase 2 (pytesseract / Google Document AI).
"""
from __future__ import annotations

import csv
import io
import logging
import math
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Mime-type → handler map ────────────────────────────────────────────────────

_SUPPORTED_MIME_TYPES: dict[str, str] = {
    "application/pdf":                                          "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword":                                       "docx",
    "text/plain":                                               "txt",
    "text/markdown":                                            "txt",
    "text/csv":                                                 "csv",
    "application/csv":                                          "csv",
}

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv"}


def is_supported(mime_type: str, filename: str) -> bool:
    """Return True if the file type is supported for ingestion."""
    ext = Path(filename).suffix.lower()
    return mime_type in _SUPPORTED_MIME_TYPES or ext in SUPPORTED_EXTENSIONS


def detect_file_type(mime_type: str, filename: str) -> str:
    """Return the canonical file_type string for a document."""
    if mime_type in _SUPPORTED_MIME_TYPES:
        return _SUPPORTED_MIME_TYPES[mime_type]
    ext = Path(filename).suffix.lower()
    return ext.lstrip(".") or "unknown"


def validate_magic_bytes(file_bytes: bytes, file_type: str) -> None:
    """Validate file content headers against dangerous binary formats and spoofing."""
    if not file_bytes:
        raise ValueError("File content is empty.")

    header = file_bytes[:16]

    # Block executable binaries regardless of declared extension
    if header.startswith(b"MZ") or header.startswith(b"\x7fELF") or header.startswith(b"\xfe\xed\xfa") or header.startswith(b"\xcf\xfa\xed\xfe"):
        raise ValueError("Security violation: Executable binary files are strictly prohibited.")

    if file_type == "pdf":
        if b"%PDF-" not in header[:10]:
            raise ValueError("Security violation: Invalid PDF header. File content does not match PDF format.")
    elif file_type == "docx":
        if not header.startswith(b"PK\x03\x04"):
            raise ValueError("Security violation: Invalid DOCX header. File content does not match DOCX (ZIP) format.")
    elif file_type in ("txt", "md", "csv"):
        # Ensure text files do not contain excessive null bytes (binary file spoofing)
        null_count = file_bytes[:1024].count(b"\x00")
        if null_count > 5:
            raise ValueError("Security violation: Binary file content detected in text/CSV upload.")


# ── Parsing ───────────────────────────────────────────────────────────────────


def parse_document(file_bytes: bytes, mime_type: str, filename: str) -> list[tuple[str, int | None]]:
    """
    Extract text from a document.

    Returns:
        list of (text_segment, page_number) tuples.
        page_number is set for PDFs; None for all other formats.

    Raises:
        ValueError  — unsupported file type or unreadable/corrupted/encrypted file.
        RuntimeError — unexpected parse failure.
    """
    file_type = detect_file_type(mime_type, filename)
    validate_magic_bytes(file_bytes, file_type)

    if file_type == "pdf":
        return _parse_pdf(file_bytes)
    elif file_type == "docx":
        return [(_parse_docx(file_bytes), None)]
    elif file_type in ("txt", "md"):
        return [(_parse_txt(file_bytes), None)]
    elif file_type == "csv":
        return [(_parse_csv(file_bytes), None)]
    else:
        raise ValueError(
            f"Unsupported file type '{file_type}' (mime: {mime_type}). "
            f"Supported: PDF, DOCX, TXT, MD, CSV."
        )


def _parse_pdf(file_bytes: bytes) -> list[tuple[str, int]]:
    """Extract text from a text-native PDF using pypdf."""
    try:
        from pypdf import PdfReader  # noqa: PLC0415
        from pypdf.errors import FileNotDecryptedError, PyPdfError  # noqa: PLC0415
        reader = PdfReader(io.BytesIO(file_bytes))

        # Check for password protection / encryption
        if reader.is_encrypted:
            try:
                decrypted = reader.decrypt("")
                if not decrypted:
                    raise ValueError("Document is password-protected or encrypted. Please remove password protection and try again.")
            except Exception as decrypt_exc:
                raise ValueError("Document is password-protected or encrypted. Please remove password protection and try again.") from decrypt_exc

        pages: list[tuple[str, int]] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append((text, page_num))
        if not pages:
            raise ValueError(
                "No extractable text found in the PDF. "
                "The document may be scanned/image-only (OCR support coming in future release)."
            )
        return pages
    except ImportError as exc:
        raise RuntimeError("pypdf is not installed. Run: uv add pypdf") from exc
    except ValueError:
        raise
    except Exception as exc:
        if "password-protected" in str(exc) or "encrypted" in str(exc) or "No extractable text" in str(exc):
            raise
        raise RuntimeError(f"PDF parsing failed: {exc}") from exc


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document  # noqa: PLC0415
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        if not paragraphs:
            raise ValueError("No extractable text found in the DOCX document.")
        return "\n\n".join(paragraphs)
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Run: uv add python-docx") from exc
    except Exception as exc:
        if "No extractable text" in str(exc):
            raise
        raise RuntimeError(f"DOCX parsing failed: {exc}") from exc


def _parse_txt(file_bytes: bytes) -> str:
    """Decode a plain-text or markdown file."""
    try:
        text = file_bytes.decode("utf-8", errors="replace").strip()
        if not text:
            raise ValueError("The text file is empty.")
        return text
    except Exception as exc:
        if "empty" in str(exc):
            raise
        raise RuntimeError(f"Text file parsing failed: {exc}") from exc


def _parse_csv(file_bytes: bytes) -> str:
    """
    Convert a CSV file to a human-readable text representation.
    Each row becomes a comma-separated "column: value" sentence suitable
    for semantic search.
    """
    try:
        text_io = io.StringIO(file_bytes.decode("utf-8", errors="replace"))
        reader = csv.DictReader(text_io)
        lines: list[str] = []
        if reader.fieldnames:
            lines.append(f"Columns: {', '.join(reader.fieldnames)}")
            lines.append("")
        for i, row in enumerate(reader):
            line = ", ".join(f"{k}: {v}" for k, v in row.items() if v is not None)
            if line:
                lines.append(line)
            if i >= 4999:  # Cap at 5000 rows to prevent runaway memory usage
                lines.append("... (truncated at 5000 rows)")
                break
        if not lines:
            raise ValueError("The CSV file is empty or has no parseable rows.")
        return "\n".join(lines)
    except Exception as exc:
        if "empty" in str(exc):
            raise
        raise RuntimeError(f"CSV parsing failed: {exc}") from exc


# ── Chunking ──────────────────────────────────────────────────────────────────


def estimate_tokens(text: str) -> int:
    """Lightweight token count estimate: ~1.3 tokens per word (no tokenizer needed)."""
    return max(1, math.ceil(len(text.split()) * 1.3))


def chunk_text(
    text: str,
    chunk_size: int | None = None,
    overlap: int | None = None,
) -> list[tuple[str, int]]:
    """
    Split text into overlapping word-window chunks, respecting paragraph breaks.

    Args:
        text:        Full text string.
        chunk_size:  Target chunk size in tokens (default: settings.DOCUMENTS_CHUNK_SIZE_TOKENS).
        overlap:     Overlap between consecutive chunks in tokens (default: settings.DOCUMENTS_CHUNK_OVERLAP_TOKENS).

    Returns:
        List of (chunk_text, estimated_token_count) tuples.
    """
    from app.config.settings import settings  # noqa: PLC0415 — lazy import to avoid circular

    chunk_size = chunk_size or settings.DOCUMENTS_CHUNK_SIZE_TOKENS
    overlap = overlap or settings.DOCUMENTS_CHUNK_OVERLAP_TOKENS

    # Convert token targets to approximate word counts (1 token ≈ 0.77 words)
    chunk_words = max(10, int(chunk_size * 0.77))
    overlap_words = max(1, int(overlap * 0.77))
    step = max(1, chunk_words - overlap_words)

    # Split by paragraphs first to respect natural boundaries, then by words
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    all_words: list[str] = []
    for para in paragraphs:
        all_words.extend(para.split())
        all_words.append("\n\n")  # paragraph separator sentinel

    # Remove trailing sentinel
    while all_words and all_words[-1] == "\n\n":
        all_words.pop()

    chunks: list[tuple[str, int]] = []
    i = 0
    while i < len(all_words):
        window = all_words[i : i + chunk_words]
        # Clean up stray sentinels inside window
        chunk_str = " ".join(w for w in window if w != "\n\n").strip()
        if chunk_str:
            tok_count = estimate_tokens(chunk_str)
            chunks.append((chunk_str, tok_count))
        i += step

    return chunks if chunks else [(text[:2000], estimate_tokens(text[:2000]))]


def deduplicate_chunks(chunks: list[tuple[str, int]]) -> list[tuple[str, int]]:
    """Remove consecutive near-duplicate chunks (identical first 80 chars)."""
    seen: set[str] = set()
    result: list[tuple[str, int]] = []
    for text, tok in chunks:
        key = text[:80].lower().strip()
        if key not in seen:
            seen.add(key)
            result.append((text, tok))
    return result


# ── Full ingestion pipeline ────────────────────────────────────────────────────


def run_ingestion_pipeline(
    document_id: str,
    user_id: str,
    file_bytes: bytes,
    mime_type: str,
    filename: str,
) -> None:
    """
    Full synchronous ingestion pipeline. Designed to run in a background thread
    (via asyncio.to_thread or FastAPI BackgroundTasks).

    Steps:
      1. Parse → text segments
      2. Chunk each segment
      3. Deduplicate
      4. Embed all chunks via the configured EmbeddingProvider
      5. Bulk-upsert chunks into document_chunks
      6. Update document_records status to 'ready'

    On any failure, sets status='error' with a user-readable message before raising.
    """
    from app.repositories.document_repository import (  # noqa: PLC0415
        update_document_status,
        insert_chunks_batch,
        delete_document_chunks,
    )
    from app.capabilities.documents.embedding import get_embedding_provider  # noqa: PLC0415
    from app.capabilities.documents.models import DocumentChunk  # noqa: PLC0415

    logger.info(
        "Starting ingestion pipeline: document_id=%s user=%s filename=%s",
        document_id, user_id, filename
    )

    try:
        # ── 1. Parse ──────────────────────────────────────────────────────────
        segments = parse_document(file_bytes, mime_type, filename)
        page_count = max((p for _, p in segments if p is not None), default=None)
        logger.info("Parsed %d segment(s) from '%s'. Max page: %s", len(segments), filename, page_count)

        # ── 2. Chunk ──────────────────────────────────────────────────────────
        raw_chunks: list[tuple[str, int, int | None]] = []  # (text, tokens, page_number)
        for text, page_num in segments:
            for chunk_text, tok in chunk_text_for_segment(text):
                raw_chunks.append((chunk_text, tok, page_num))

        raw_chunks = [(t, tk, p) for t, tk, p in raw_chunks]  # identity pass (type safety)
        logger.info("Produced %d raw chunk(s) before deduplication.", len(raw_chunks))

        # Deduplicate
        seen_keys: set[str] = set()
        deduped: list[tuple[str, int, int | None]] = []
        for text, tok, page_num in raw_chunks:
            key = text[:80].lower().strip()
            if key not in seen_keys:
                seen_keys.add(key)
                deduped.append((text, tok, page_num))

        logger.info("%d chunk(s) after deduplication.", len(deduped))

        # ── 3. Embed ──────────────────────────────────────────────────────────
        provider = get_embedding_provider()
        texts = [t for t, _, _ in deduped]
        embeddings = provider.embed_documents(texts)
        logger.info("Embedded %d chunk(s) using %s.", len(embeddings), provider.model_name)

        # ── 4. Assemble DocumentChunk objects ─────────────────────────────────
        # Clean out any previously stored chunks for this doc (handles retry on crash)
        delete_document_chunks(document_id)

        chunk_objects: list[DocumentChunk] = []
        for idx, ((text, tok, page_num), embedding) in enumerate(zip(deduped, embeddings)):
            chunk_objects.append(DocumentChunk(
                document_id=document_id,
                user_id=user_id,
                chunk_index=idx,
                content=text,
                token_count=tok,
                page_number=page_num,
                embedding=embedding,
                embedding_model=provider.model_name,
            ))

        # ── 5. Bulk-upsert ────────────────────────────────────────────────────
        insert_chunks_batch(chunk_objects)

        # ── 6. Mark document as ready ─────────────────────────────────────────
        update_document_status(
            document_id=document_id,
            status="ready",
            page_count=page_count,
            chunk_count=len(chunk_objects),
            embedding_model=provider.model_name,
        )
        logger.info(
            "Ingestion complete for document_id=%s: %d chunk(s) stored.",
            document_id, len(chunk_objects)
        )

    except Exception as exc:  # noqa: BLE001
        logger.exception("Ingestion pipeline failed for document_id=%s", document_id)
        update_document_status(
            document_id=document_id,
            status="error",
            error_message=str(exc)[:500],  # cap length for DB storage
        )


def chunk_text_for_segment(text: str) -> list[tuple[str, int]]:
    """Thin wrapper: chunk a single text segment and return (text, token_count) pairs."""
    return chunk_text(text)
